# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Розділ_2_1_архітектура_DinoJelly.docx"


PARAGRAPHS = [
    "Архітектурно проєкт DinoJelly належить до класу сучасних монолітних веборієнтованих інформаційних систем. Основна серверна логіка, маршрутизація, контроль доступу, робота з базою даних, доменні сервіси, фонова обробка та адміністративні функції зосереджені в межах одного Laravel-застосунку. Водночас клієнтський інтерфейс реалізовано на Vue 3 із використанням Inertia.js, що дозволяє поєднати серверно-керовану логіку з реактивною поведінкою сторінок без побудови окремого REST або GraphQL API для кожного сценарію взаємодії.",
    "Загальну схему розгортання системи наведено на рисунку 2.1. Діаграма показує, що користувач взаємодіє із системою через браузер за протоколом HTTPS. Запити надходять до NGINX Ingress, після чого передаються до сервісу laravel-service і далі до розгортання laravel-app. Застосунок виконується в Kubernetes-кластері та має три репліки, що підвищує доступність системи й дозволяє обробляти кілька паралельних запитів. Для збереження конфігураційних параметрів використовується laravel-secrets, де розміщуються ключ застосунку та облікові дані доступу до бази даних.",
    "Окремими вузлами розгортання виступають MySQL і Redis. MySQL використовується як основне сховище даних для користувачів, товарів, замовлень, платежів, складських залишків та інших сутностей предметної області. Для збереження даних MySQL застосовується постійне сховище mysql-pvc, підключене до контейнера бази даних. Redis використовується для кешування, збереження сесій і роботи черг, що важливо для фонових процесів, зокрема обробки повторних задач і надсилання повідомлень.",
    "Компонентну структуру програмної системи наведено на рисунку 2.2. Вона відображає логічний поділ DinoJelly на кілька взаємопов’язаних частин. Рівень представлення містить клієнтський інтерфейс Client UI, побудований на Vue 3 та Inertia.js. Він взаємодіє з Routing Layer, який відповідає за маршрутизацію вебзапитів, застосування middleware, перевірку авторизації та захист від CSRF-атак.",
    "Наступним рівнем є Controllers, що об’єднує користувацькі й адміністративні контролери. Вони приймають HTTP-запити, виконують первинну валідацію, визначають необхідний сценарій роботи та делегують бізнес-операції доменним сервісам. Такий підхід зменшує навантаження на контролери й дозволяє зберігати основну логіку обробки замовлень у спеціалізованих сервісах.",
    "Доменний рівень представлений компонентом Domain Services. До нього належать CatalogService, CheckoutService, OrderService, PaymentService, InventoryService та інші сервіси, які реалізують ключові бізнес-правила системи. Саме цей рівень відповідає за формування замовлення, перевірку складських залишків, створення платежу, оновлення статусів, облік подій замовлення та взаємодію з іншими модулями системи.",
    "Окремо виділено Marketing Services, які забезпечують роботу реферальної програми, подарункових сертифікатів, рекомендаційного помічника та механізму відновлення покинутих кошиків. Jobs and Commands відповідають за запуск фонових задач і команд планувальника. Infrastructure Adapters забезпечують технічну взаємодію із зовнішніми службами та інфраструктурою: базою даних, Redis, поштовим транспортом і webhook-обробниками платежів. Persistence Model містить Eloquent-моделі, ресурси та міграції, через які доменні сервіси працюють із даними.",
    "Для системи обліку замовлень така архітектура є доцільною, оскільки життєвий цикл замовлення охоплює багато взаємопов’язаних операцій: перевірку кошика, резервування товарів, вибір адреси й слота доставки, створення замовлення, формування платежу, зміну статусів і фіксацію подій. Винесення цих операцій у доменні сервіси підвищує цілісність логіки, полегшує тестування та спрощує подальше розширення системи.",
    "Важливою особливістю DinoJelly є те, що користувацька та адміністративна частини працюють з одними й тими самими сутностями: товарами, замовленнями, платежами, користувачами, залишками та відгуками. Завдяки цьому система не дублює бізнес-логіку в різних частинах застосунку, а використовує спільний набір моделей і сервісів. Користувач послідовно переходить від перегляду каталогу до кошика, оформлення замовлення та перегляду історії покупок, тоді як адміністратор працює з каталогом, складом, платежами, замовленнями й аналітичними показниками через окремий адміністративний інтерфейс.",
]


TABLE_ROWS = [
    ("Представлення", "Vue 3 сторінки, компоненти, Inertia.js", "Відображення каталогу, кошика, checkout, історії замовлень та адміністративного інтерфейсу"),
    ("Маршрутизація та HTTP-рівень", "routes/web.php, middleware, контролери", "Приймання HTTP-запитів, перевірка доступу, маршрутизація користувацьких і адміністративних сценаріїв"),
    ("Доменний рівень", "CheckoutService, OrderService, PaymentService, InventoryService, CatalogService", "Реалізація бізнес-правил обробки замовлень, платежів, складу та каталогу"),
    ("Маркетинговий рівень", "ReferralService, GiftCardService, RecommendationAssistantService, AbandonedCartService", "Підтримка реферальної програми, подарункових сертифікатів, рекомендацій і відновлення покинутих кошиків"),
    ("Рівень даних", "Eloquent-моделі, API resources, міграції MySQL", "Збереження та зв’язування сутностей системи"),
    ("Інфраструктурний рівень", "Kubernetes, NGINX Ingress, MySQL, Redis, PVC, secrets", "Розгортання, маршрутизація трафіку, збереження даних, кешування, сесії та черги"),
]


def set_run_font(run, size: int = 14, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_font(paragraph, size: int = 14, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size, bold)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_centered_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.line_spacing = 1.5
    run = title.add_run("2.1 Архітектурна концепція та модульна структура")
    set_run_font(run, bold=True)

    for index, text in enumerate(PARAGRAPHS):
        add_paragraph(doc, text)
        if index == 2:
            add_centered_caption(doc, "Рисунок 2.1 – Діаграма розгортання системи DinoJelly")
        if index == 6:
            add_centered_caption(doc, "Рисунок 2.2 – Діаграма компонентів програмної системи")

    add_centered_caption(doc, "Таблиця 2.1 – Архітектурні рівні системи")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headers = ["Рівень", "Основні елементи", "Призначення"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, bold=True)

    for row_data in TABLE_ROWS:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_data):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=12)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
